"""Automação: ler atividades de um PDF aberto no navegador e escrever no Word **ou** gerar arquivo.

Fluxo:
- Foca janela do PDF (Chrome/Edge), faz OCR da tela e rola (PageDown) até estabilizar.
- Extrai blocos "Atividade X" e perguntas.
- Gera um texto organizado com seções e respostas sugeridas (quando reconhecível).
- Saída:
    - output_mode="word": foca janela do Word e digita o conteúdo.
    - output_mode="docx": gera um arquivo .docx.
    - output_mode="pdf": gera um arquivo .pdf.

Limitações:
- OCR é best-effort e depende de zoom/contraste.
- Para máxima precisão, deixe o PDF em 100%-125% e em tela cheia.

Segurança:
- Pressiona teclas / clica em janelas => HIGH (deve passar por HITL).
"""

from __future__ import annotations

import re
import time
from dataclasses import dataclass
from pathlib import Path
import textwrap
import os
import shutil
import sys
import subprocess
from typing import Any

from omniscia.core.tools import ToolRegistry, ToolSpec
from omniscia.core.types import ToolResult
from omniscia.modules.os_control.filesystem import resolve_known_folder_prefixed_path


def register_edu_tools(registry: ToolRegistry) -> None:
    registry.register(
        ToolSpec(
            name="edu.pdf_word_autofill",
            description=(
                "Lê atividades de um PDF aberto (OCR + rolagem) e envia a saída para Word ou arquivo. "
                "Args: pdf_title_contains?, assume_focused_pdf?, output_mode(word|docx|pdf)?, word_title_contains?, out_path?, "
                "overwrite?, solve_with_llm?, llm_max_questions?, max_scrolls?, duration_s?, settle_ms?"
            ),
            risk="HIGH",
            fn=_pdf_word_autofill,
        )
    )


def _require_pyautogui():
    try:
        import pyautogui

        pyautogui.FAILSAFE = True
        pyautogui.PAUSE = 0.02
        return pyautogui, None
    except Exception as exc:  # noqa: BLE001
        return None, f"pyautogui indisponível: {exc}"


def _require_vision_ocr():
    try:
        import mss
        from PIL import Image
        import pytesseract

        # Configura tesseract_cmd se possível.
        # Preferência: OMNI_TESSERACT_CMD (Settings) -> PATH -> caminhos comuns no Windows.
        try:
            from omniscia.core.config import Settings

            s = Settings.load()
            tcmd = (getattr(s, "tesseract_cmd", None) or "").strip() or None
        except Exception:
            tcmd = None

        if tcmd and Path(tcmd).exists():
            pytesseract.pytesseract.tesseract_cmd = tcmd
        else:
            which = shutil.which("tesseract")
            if which:
                pytesseract.pytesseract.tesseract_cmd = which
            else:
                candidates: list[str] = []
                pf = os.environ.get("ProgramFiles")
                pfx86 = os.environ.get("ProgramFiles(x86)")
                local = os.environ.get("LOCALAPPDATA")
                if pf:
                    candidates.append(str(Path(pf) / "Tesseract-OCR" / "tesseract.exe"))
                if pfx86:
                    candidates.append(str(Path(pfx86) / "Tesseract-OCR" / "tesseract.exe"))
                if local:
                    candidates.append(str(Path(local) / "Programs" / "Tesseract-OCR" / "tesseract.exe"))

                for c in candidates:
                    if Path(c).exists():
                        pytesseract.pytesseract.tesseract_cmd = c
                        break

        return mss, Image, pytesseract, None
    except Exception as exc:  # noqa: BLE001
        return None, None, None, f"deps OCR ausentes (mss/pillow/pytesseract): {exc}"


def _focus_window_best_effort(pyautogui, title_contains: str, *, timeout_s: float = 2.5) -> bool:
    title_contains = (title_contains or "").strip()
    if not title_contains:
        return False
    try:
        from omniscia.modules.os_control.win_windows import focus_window_by_title_contains

        rect = focus_window_by_title_contains(title_contains, timeout_s=timeout_s)
        if rect:
            cx = int((rect["left"] + rect["right"]) // 2)
            cy = int((rect["top"] + rect["bottom"]) // 2)
            pyautogui.click(x=cx, y=cy, button="left")
            return True
    except Exception:
        return False
    return False


def _focus_word_best_effort(pyautogui, title_hint: str, *, timeout_s: float = 6.0) -> bool:
    """Best-effort focus for Microsoft Word.

    Word titles are often localized/dynamic, so we try:
    - user-provided hint
    - common title fragments (pt-BR/en)
    - Win32 class name fallback (OpusApp)
    """

    hint = (title_hint or "").strip() or "Word"
    candidates = [hint, "Microsoft Word", "Word", "Documento", "Document", "Sem Título", "Untitled"]
    seen: set[str] = set()
    cand2: list[str] = []
    for c in candidates:
        c2 = (c or "").strip()
        if not c2:
            continue
        key = c2.casefold()
        if key in seen:
            continue
        seen.add(key)
        cand2.append(c2)

    # Try by title fragments within a deadline.
    deadline = time.time() + float(timeout_s)
    while time.time() < deadline:
        for t in cand2:
            if _focus_window_best_effort(pyautogui, t, timeout_s=0.9):
                return True
        time.sleep(0.1)

    # Fallback: focus by class name.
    try:
        from omniscia.modules.os_control.win_windows import focus_window_by_class_name

        rect = focus_window_by_class_name("OpusApp", timeout_s=2.5)
        if rect:
            cx = int((rect["left"] + rect["right"]) // 2)
            cy = int((rect["top"] + rect["bottom"]) // 2)
            pyautogui.click(x=cx, y=cy, button="left")
            return True
    except Exception:
        return False

    return False


def _try_launch_word_best_effort() -> None:
    if not sys.platform.startswith("win"):
        return
    try:
        # Prefer shell start to rely on file associations / PATH.
        subprocess.Popen(["cmd", "/c", "start", "", "winword"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except Exception:
        # Best-effort only.
        return


def _grab_screen_gray(mss_mod, pil_image_mod):
    with mss_mod.mss() as sct:
        monitor = sct.monitors[1]
        shot = sct.grab(monitor)
        img = pil_image_mod.frombytes("RGB", shot.size, shot.rgb)
    return img.convert("L")


def _ocr_text(pytesseract_mod, gray_img) -> str:
    # Normalizações simples para reduzir ruído.
    try:
        text = pytesseract_mod.image_to_string(gray_img, lang="por")
    except Exception:
        # Fallback: sem especificar idioma.
        text = pytesseract_mod.image_to_string(gray_img)
    text = text.replace("\r", "")
    # Colapsa múltiplos espaços.
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _ensure_word_blank_document(pyautogui) -> None:
    # Ctrl+N cria um novo documento em branco mesmo quando o Word abre na tela inicial.
    try:
        pyautogui.hotkey("ctrl", "n")
        time.sleep(0.85)
    except Exception:
        pass


@dataclass(frozen=True)
class ExtractedDoc:
    raw_pages: list[str]
    merged_text: str


def _collect_pdf_text_by_scrolling(
    *,
    pyautogui,
    mss_mod,
    pil_image_mod,
    pytesseract_mod,
    max_scrolls: int,
    duration_s: float,
) -> ExtractedDoc:
    pages: list[str] = []
    seen_fingerprints: dict[str, int] = {}

    start = time.time()
    for _i in range(max_scrolls):
        if time.time() - start >= duration_s:
            break

        gray = _grab_screen_gray(mss_mod, pil_image_mod)
        text = _ocr_text(pytesseract_mod, gray)
        fp = re.sub(r"\W+", "", text.lower())[:600]
        if fp:
            seen_fingerprints[fp] = seen_fingerprints.get(fp, 0) + 1

        if text:
            pages.append(text)

        # Se estamos vendo a mesma página repetidas vezes, provavelmente chegamos no fim.
        if fp and seen_fingerprints.get(fp, 0) >= 3:
            break

        # Rola
        pyautogui.press("pagedown")
        time.sleep(0.20)

    merged = "\n\n".join(pages)
    return ExtractedDoc(raw_pages=pages, merged_text=merged)


def _split_activities(text: str) -> list[tuple[str, str]]:
    """Retorna lista de (titulo, corpo)."""

    if not text:
        return []

    # Normaliza variações: "Atividade 1" / "Atividade 1 (prática)".
    pat = re.compile(r"(?im)^\s*atividade\s*(\d+)\b.*$")
    matches = list(pat.finditer(text))
    if not matches:
        return []

    blocks: list[tuple[str, str]] = []
    for idx, m in enumerate(matches):
        start = m.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        title_line = (m.group(0) or "").strip()
        body = text[start:end].strip()
        blocks.append((title_line, body))
    return blocks


def _extract_questions(activity_block: str) -> list[str]:
    """Extrai perguntas/bullets mais prováveis."""

    lines = [ln.strip() for ln in (activity_block or "").split("\n") if ln.strip()]
    out: list[str] = []

    for ln in lines:
        # bullets comuns
        if ln.startswith(("•", "-", "*")):
            out.append(ln.lstrip("•-* ").strip())
            continue
        # também aceita linhas longas que parecem enunciado.
        if len(ln) >= 40 and ("?" in ln or re.search(r"\b(explique|utilizando|finalize|abra)\b", ln, re.I)):
            out.append(ln)

    # de-dup simples
    dedup: list[str] = []
    seen: set[str] = set()
    for q in out:
        key = re.sub(r"\W+", "", q.lower())[:160]
        if key and key not in seen:
            seen.add(key)
            dedup.append(q)
    return dedup


def _suggest_answer(question: str) -> str:
    qn = (question or "").lower()

    # Powershell: notepad open + stop by name
    if "powershell" in qn and ("bloco de notas" in qn or "notepad" in qn) and ("finalize" in qn or "feche" in qn):
        return (
            "Comandos (PowerShell):\n"
            "1) Start-Process notepad.exe\n"
            "2) Stop-Process -Name notepad\n\n"
            "Parâmetros:\n"
            "- Start-Process: inicia um executável/programa.\n"
            "- Stop-Process -Name: encerra processos pelo nome (sem .exe)."
        )

    # Powershell: calculator open + stop by PID
    if "powershell" in qn and ("calculadora" in qn or "calc" in qn) and ("numero de processo" in qn or "pid" in qn):
        return (
            "Comandos (PowerShell):\n"
            "1) $p = Start-Process calc.exe -PassThru\n"
            "2) Stop-Process -Id $p.Id\n\n"
            "Parâmetros:\n"
            "- Start-Process -PassThru: retorna o objeto do processo (com Id/PID).\n"
            "- Stop-Process -Id: encerra pelo PID."
        )

    # Ajuda -?
    if re.search(r"\b(-\?)\b", qn) or "parametro" in qn and "-?" in qn:
        return (
            "No PowerShell, `-?` mostra a ajuda do comando, incluindo sintaxe e parâmetros.\n"
            "Ex.: Get-Help Start-Process -Full  (equivalente/mais completo)."
        )

    return "(Resposta sugerida: revisar e completar conforme o enunciado.)"


def _llm_can_run() -> tuple[bool, str | None, Any | None]:
    """Verifica se há config/deps para usar LLM. Retorna (ok, err, settings)."""

    try:
        from omniscia.core.config import Settings
        from omniscia.core.litellm_env import apply_litellm_env, provider_requires_api_key
    except Exception as exc:  # noqa: BLE001
        return False, f"Core LLM indisponível: {exc}", None

    settings = Settings.load()
    provider = (settings.llm_provider or "").strip()
    model = (settings.llm_model or "").strip()
    if not provider or not model:
        return False, "LLM não configurado (defina OMNI_LLM_PROVIDER e OMNI_LLM_MODEL)", settings

    needs_key = provider_requires_api_key(provider)
    has_key = bool((settings.llm_api_key or "").strip())
    if needs_key and not has_key:
        return False, "LLM requer chave (defina OMNI_LLM_API_KEY)", settings

    try:
        import litellm  # noqa: F401
        from litellm import completion  # noqa: F401
    except Exception as exc:  # noqa: BLE001
        return False, f"litellm não disponível: {exc}", settings

    apply_litellm_env(settings)
    return True, None, settings


def _answer_with_llm(*, settings: Any, activity_title: str, activity_body: str, question: str) -> str:
    """Gera resposta via LLM. Se falhar, levanta exceção para fallback."""

    from litellm import completion

    body = (activity_body or "").strip()
    if len(body) > 2800:
        body = body[:2800] + "\n..."

    system = (
        "Você é um assistente educacional. Responda em PT-BR, direto ao ponto, "
        "com passos/comandos quando aplicável. Não invente dados; se faltar info no enunciado, "
        "explique o que está faltando e dê a melhor orientação possível."
    )

    user = (
        "Contexto extraído por OCR (pode conter ruído).\n\n"
        f"{activity_title.strip()}\n\n"
        f"Trecho do enunciado:\n{body}\n\n"
        f"Pergunta: {question.strip()}\n\n"
        "Responda com uma solução completa e organizada."
    )

    resp_obj: Any = completion(
        model=str(getattr(settings, "llm_model", "") or "").strip(),
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        temperature=0.2,
        max_tokens=600,
    )

    content = ""
    try:
        if isinstance(resp_obj, dict):
            content = resp_obj.get("choices", [{}])[0].get("message", {}).get("content", "")
        else:
            choices = getattr(resp_obj, "choices", None)
            if choices:
                msg = getattr(choices[0], "message", None)
                if isinstance(msg, dict):
                    content = msg.get("content", "")
                else:
                    content = getattr(msg, "content", "") or ""
    except Exception:
        content = ""
    return (content or "").strip() or "(LLM retornou vazio; usar fallback.)"


def _format_document(
    extracted: ExtractedDoc,
    *,
    solve_with_llm: bool = False,
    llm_max_questions: int = 12,
) -> tuple[str, dict[str, Any]]:
    blocks = _split_activities(extracted.merged_text)
    if not blocks:
        # fallback: dump do OCR
        return (
            "ATIVIDADES (OCR)\n\n"
            "Não consegui segmentar por 'Atividade X'. Segue o texto OCR bruto:\n\n"
            + extracted.merged_text.strip()
        ), {"llm_used": False, "llm_error": None, "llm_questions": 0}

    llm_used = False
    llm_error: str | None = None
    llm_questions = 0
    llm_settings: Any | None = None
    if solve_with_llm:
        ok, err, settings = _llm_can_run()
        if ok:
            llm_settings = settings
        else:
            llm_error = err
            solve_with_llm = False

    parts: list[str] = ["ATIVIDADES — preenchimento no Word\n"]
    for title, body in blocks:
        parts.append(title)
        qs = _extract_questions(body)
        if not qs:
            parts.append("(Não consegui extrair perguntas com confiança; revisar OCR.)\n")
            continue
        for i, q in enumerate(qs, start=1):
            parts.append(f"{i}. {q}")
            parts.append("Resposta:")
            if solve_with_llm and llm_settings is not None and llm_questions < llm_max_questions:
                try:
                    ans = _answer_with_llm(
                        settings=llm_settings,
                        activity_title=title,
                        activity_body=body,
                        question=q,
                    )
                    llm_used = True
                    llm_questions += 1
                except Exception as exc:  # noqa: BLE001
                    # Se bater rate limit, desliga LLM pro restante para evitar cascata.
                    msg_err = str(exc)
                    llm_error = f"falha LLM: {msg_err}"
                    if re.search(r"rate\s*limit|429|too\s*many\s*requests|tpm|tokens\s*per\s*minute", msg_err, re.I):
                        solve_with_llm = False
                    ans = _suggest_answer(q)
                parts.append(ans)
            else:
                parts.append(_suggest_answer(q))
            parts.append("")
        parts.append("")

    return (
        "\n".join(parts).strip() + "\n",
        {"llm_used": llm_used, "llm_error": llm_error, "llm_questions": llm_questions},
    )


def _write_audit_text(content: str) -> None:
    try:
        out = Path("data/tmp/atividades_word.txt")
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(content, encoding="utf-8")
    except Exception:
        pass


def _write_docx(out_path: Path, content: str) -> tuple[bool, str | None]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # noqa: BLE001
        return (
            False,
            "python-docx não está instalado. Instale com: pip install -e \".[docs]\"  (ou pip install python-docx). "
            + f"Detalhe: {exc}",
        )

    doc = Document()
    for line in content.splitlines():
        ln = (line or "").rstrip("\n")
        if not ln.strip():
            doc.add_paragraph("")
            continue

        if re.match(r"(?i)^atividades\b", ln.strip()):
            doc.add_heading(ln.strip(), level=0)
            continue
        if re.match(r"(?i)^atividade\s*\d+\b", ln.strip()):
            doc.add_heading(ln.strip(), level=1)
            continue

        doc.add_paragraph(ln)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return True, None


def _write_pdf(out_path: Path, content: str) -> tuple[bool, str | None]:
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception as exc:  # noqa: BLE001
        return (
            False,
            "reportlab não está instalado. Instale com: pip install -e \".[docs]\"  (ou pip install reportlab). "
            + f"Detalhe: {exc}",
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)

    c = canvas.Canvas(str(out_path), pagesize=A4)
    width, height = A4
    margin = 36
    y = height - margin
    line_h = 14

    for para in content.splitlines():
        if not para.strip():
            y -= line_h
            continue

        for ln in textwrap.wrap(para, width=100, break_long_words=False, break_on_hyphens=False):
            if y < margin:
                c.showPage()
                y = height - margin
            c.drawString(margin, y, ln)
            y -= line_h

    c.save()
    return True, None


def _pdf_word_autofill(args: dict[str, Any]) -> ToolResult:
    pyautogui, err = _require_pyautogui()
    if pyautogui is None:
        return ToolResult(status="error", error=err)

    mss_mod, pil_image_mod, pytesseract_mod, oerr = _require_vision_ocr()
    if mss_mod is None:
        return ToolResult(status="error", error=oerr)

    pdf_title = str(args.get("pdf_title_contains", "") or "").strip()
    assume_focused_pdf = bool(args.get("assume_focused_pdf", False))

    output_mode = str(args.get("output_mode", "word") or "word").strip().lower()
    if output_mode not in {"word", "docx", "pdf"}:
        return ToolResult(status="error", error="output_mode inválido (use: word|docx|pdf)")

    word_title = str(args.get("word_title_contains", "") or "").strip() or "Word"
    out_path_raw = str(args.get("out_path", "") or "").strip().replace("\\", "/")
    overwrite = bool(args.get("overwrite", True))

    solve_with_llm = bool(args.get("solve_with_llm", False))
    llm_max_questions = int(args.get("llm_max_questions", 12) or 12)
    llm_max_questions = max(1, min(80, llm_max_questions))

    settle_ms = int(args.get("settle_ms", 650) or 650)
    max_scrolls = int(args.get("max_scrolls", 18) or 18)
    duration_s = float(args.get("duration_s", 35.0) or 35.0)

    if not pdf_title and not assume_focused_pdf:
        return ToolResult(
            status="error",
            error=(
                "pdf_title_contains é obrigatório (ex: 'Atividades.pdf') "
                "ou defina assume_focused_pdf=true (você coloca a janela do PDF em foco antes)."
            ),
        )

    max_scrolls = max(3, min(60, max_scrolls))
    duration_s = max(5.0, min(180.0, duration_s))
    settle_ms = max(0, min(5000, settle_ms))

    # 1) Foca PDF (ou assume que já está em foco)
    if assume_focused_pdf:
        time.sleep(settle_ms / 1000.0)
    else:
        ok_pdf = _focus_window_best_effort(pyautogui, pdf_title, timeout_s=4.0)
        time.sleep(settle_ms / 1000.0)
        if not ok_pdf:
            return ToolResult(status="error", error=f"Não consegui focar a janela do PDF com title_contains='{pdf_title}'")

    # 2) Coleta OCR rolando
    try:
        extracted = _collect_pdf_text_by_scrolling(
            pyautogui=pyautogui,
            mss_mod=mss_mod,
            pil_image_mod=pil_image_mod,
            pytesseract_mod=pytesseract_mod,
            max_scrolls=max_scrolls,
            duration_s=duration_s,
        )
    except getattr(pytesseract_mod, "TesseractNotFoundError", Exception):
        return ToolResult(
            status="error",
            error=(
                "Tesseract não encontrado. Instale o Tesseract (Windows) e/ou configure OMNI_TESSERACT_CMD. "
                "Caminho comum: C:/Program Files/Tesseract-OCR/tesseract.exe"
            ),
        )
    except FileNotFoundError:
        return ToolResult(
            status="error",
            error=(
                "Falha ao executar o OCR: arquivo não encontrado. "
                "Isso normalmente indica que o tesseract.exe não está instalado/no PATH. "
                "Instale o Tesseract e/ou configure OMNI_TESSERACT_CMD."
            ),
        )

    if not extracted.merged_text.strip():
        return ToolResult(status="error", error="OCR retornou vazio. Ajuste zoom/contraste e tente novamente.")

    doc_text, meta = _format_document(extracted, solve_with_llm=solve_with_llm, llm_max_questions=llm_max_questions)

    # Sempre salva cópia no workspace (auditoria/debug)
    _write_audit_text(doc_text)

    if output_mode == "word":
        # 3) Foca Word
        ok_word = _focus_word_best_effort(pyautogui, word_title, timeout_s=6.5)
        time.sleep(0.40)
        if not ok_word:
            # Fallback extra: tentar abrir o Word e focar novamente.
            _try_launch_word_best_effort()
            time.sleep(1.0)
            ok_word2 = _focus_word_best_effort(pyautogui, word_title, timeout_s=6.5)
            time.sleep(0.40)
            if not ok_word2:
                return ToolResult(status="error", error=f"Não consegui focar o Word (tente abrir um documento e deixar o Word visível). title_hint='{word_title}'")

        _ensure_word_blank_document(pyautogui)

        # 4) Digita no Word (append no cursor)
        try:
            pyautogui.write(doc_text, interval=0.003)
        except Exception as exc:  # noqa: BLE001
            return ToolResult(status="error", error=f"Falha digitando no Word: {exc}")

        note = ""
        if solve_with_llm and not meta.get("llm_used") and meta.get("llm_error"):
            note = f" (LLM indisponível: {meta.get('llm_error')})"
        if meta.get("llm_used"):
            note = f" (LLM: respondeu {meta.get('llm_questions', 0)} pergunta(s))"

        return ToolResult(
            status="ok",
            output=(
                f"OK. Extraí {len(extracted.raw_pages)} telas via OCR e escrevi no Word. "
                "Também salvei data/tmp/atividades_word.txt"
                + note
            ),
        )

    # output_mode: docx/pdf => gerar arquivo
    if not out_path_raw:
        out_path_raw = "data/tmp/atividades.docx" if output_mode == "docx" else "data/tmp/atividades.pdf"

    out_path_display = out_path_raw
    # Permitimos known folders via prefixo (ex: desktop:/atividades.docx).
    low = out_path_raw.lower()
    if low.startswith(("desktop:/", "downloads:/", "documents:/")):
        try:
            out_path = resolve_known_folder_prefixed_path(out_path_raw)
        except Exception as exc:  # noqa: BLE001
            return ToolResult(status="error", error=f"out_path inválido (known folder): {exc}")
    else:
        # Mantém restrição: somente paths relativos ao workspace (sem drive/raiz).
        if out_path_raw.startswith("/") or ":" in out_path_raw:
            return ToolResult(status="error", error="out_path inválido (use path relativo ao workspace ou desktop:/...)")
        out_path = Path(out_path_raw)

    if out_path.exists() and not overwrite:
        return ToolResult(status="error", error=f"Arquivo já existe e overwrite=false: {out_path_raw}")

    if output_mode == "docx":
        ok, werr = _write_docx(out_path, doc_text)
        if not ok:
            return ToolResult(status="error", error=werr)
    elif output_mode == "pdf":
        ok, werr = _write_pdf(out_path, doc_text)
        if not ok:
            return ToolResult(status="error", error=werr)

    note = ""
    if solve_with_llm and not meta.get("llm_used") and meta.get("llm_error"):
        note = f" (LLM indisponível: {meta.get('llm_error')})"
    if meta.get("llm_used"):
        note = f" (LLM: respondeu {meta.get('llm_questions', 0)} pergunta(s))"

    return ToolResult(
        status="ok",
        output=(
            f"OK. Extraí {len(extracted.raw_pages)} telas via OCR e gerei {out_path_display} (salvo em {str(out_path)}). "
            "Também salvei data/tmp/atividades_word.txt"
            + note
        ),
    )
